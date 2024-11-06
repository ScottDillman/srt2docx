# -*- coding: utf-8 -*-
"""Script to convert SRT files to DOCX files with tables

Example of end state:
- https://docs.google.com/document/d/1MqnG5MWRbjZBke9ja5CN-rJGijJnWBrRJa6lLOxMvRQ/edit

SRT format docs
- https://en.wikipedia.org/wiki/SubRip

DOCX docs
- https://python-docx.readthedocs.io/en/latest/

Example:

        $ python srt2docx.py


@Author: Scott Dillman <scott@bitwise.ninja>
@Date: 2024-06-25 22:47
"""

import uuid
from munch import munchify
import yaml
from loguru import logger
from pathlib import Path
from datetime import datetime,timedelta
import srt
import pytz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import inspect
import os
from docx.enum.style import WD_STYLE_TYPE
import time

def round_to_secs(dt: datetime) -> datetime:
    """round to nearest second"""
    return timedelta(seconds=int(dt.total_seconds()))

def age(x):
    return  time.time() - os.path.getmtime(x)

def getbytes(x):
    return int(os.path.getsize(x))

def readFiles(values) -> list:
    """read files from cwd"""
    logger.info("Glob in effect is: [{}]".format(values.settings.filetypes.glob))
    logger.info("Sorting in effect is: [{}][{}]".format(values.settings.sort.type,values.settings.sort.direction))

    key = None;
    reverse = False

    if values.settings.sort.direction == 'descending':
        reverse = True
    else:
        reverse = False

    ## sort files here based on settings
    match values.settings.sort.type:
        case "name":
            key = os.path.normpath
        case "size":
            key = getbytes
        case "age":
            key = age
        case _:
            key = os.path.normpath

    files = sorted(list(Path().glob(values.settings.filetypes.glob)),key=key, reverse=reverse)

    return files


def createDocument(values) -> Document:
    """Create a new docx document"""
    document = Document()

    ## set normal style to font in settings
    style = document.styles["Normal"]
    font = style.font
    font.name = values.settings.layout.fonts.Normal

    ## Add stylized heading
    styles = document.styles
    new_heading_style = styles.add_style("New Heading", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Heading 1"]
    font = new_heading_style.font
    font.name = values.settings.layout.fonts.Title
    font.size = Pt(values.settings.layout.title.size)
    font.color.rgb = RGBColor.from_string(values.settings.layout.title.color)

    return document


def closeDocument(values, name, version, document) -> None:
    """End and save docx document"""

    ## set margins
    document.sections[0].left_margin = Inches(values.settings.layout.margin_left)
    document.sections[0].right_margin = Inches(values.settings.layout.margin_right)

    ## add core properties
    cprops = document.core_properties

    ## add basic document properties from the config yaml
    cprops.author = values.settings.meta.author
    cprops.category = values.settings.meta.category
    cprops.comments = values.settings.meta.comments
    cprops.content_status = values.settings.meta.content_status
    cprops.keywords = values.settings.meta.keywords
    cprops.language = values.settings.meta.language
    cprops.subject = values.settings.meta.subject
    cprops.version = values.settings.meta.version
    cprops.last_modified_by = values.settings.meta.author

    ## UTC is expected, localization happens on the client
    tz = pytz.timezone("UTC")
    cprops.created = datetime.now(tz)
    cprops.modified = datetime.now(tz)

    ## assign a UUID to make this document unique and to tag it with eh str2docx version
    ## we put the versioon here in case we need to debug a broken file
    cprops.identifier = "v{}-{}".format(version, str(uuid.uuid4()))
    logger.info("Document unique id: [{}]", cprops.identifier)

    ## let's add a fun footer
    if values.settings.footer.show:
        section = document.sections[0]
        section.footer_distance = Inches(0.2)
        footer = section.footer
        footer_para = footer.paragraphs[0]
        logo_run = footer_para.add_run()

        ## if you want other watermarks put them in the assets directory
        ## because that is where we look for this file
        p = os.path.join(
            os.path.dirname(os.path.realpath(inspect.stack()[0][1])), "assets"
        )
        logo_run.add_picture(
            os.path.join(p, values.settings.footer.watermark),
            width=Inches(values.settings.footer.width_in),
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ## save the doc out using the base name from read file
    logger.info("Saving: [{}.docx]".format(name))
    document.save("{}.docx".format(name))


def buildTable(values, subs, title, document) -> None:
    """build the docx table and add it to the document"""
    document.add_paragraph(title, style="New Heading")

    ## add SRT data in table
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    table.allow_autofit = False

    ## set column widths in table
    count = 0
    for c in values.settings.layout.table_cols:
        table.columns[count].width = Inches(c)
        table.rows[0].cells[count].width = Inches(c)
        count = count + 1

    ## create table header
    hdr_cells = table.rows[0].cells
    count = 0
    for h in values.settings.table.headers:
        ##hdr_cells[count].text = h
        hdr_cells[count].paragraphs[0].add_run(h).bold = True
        count = count + 1

    ## iterate over the subs and add data to table
    for item in subs:
        row_cells = table.add_row().cells

        ## TODO: fixme
        ## drop microseconds
        start = round_to_secs(item.start)
        end   = round_to_secs(item.end)
        row_cells[0].text = str(start)
        row_cells[1].text = str(end)


        row_cells[2].text = str(round_to_secs(end - start))
        row_cells[3].text = "{}".format(item.content)


def processFiles(values, files: list, version: str) -> None:
    """transform files into tables"""

    d = None
    if values.settings.single_file:
        d = createDocument(values)

    ## iterate over file list
    for i in files:
        logger.info("Processing: [{}]".format(i))

        if not values.settings.single_file:
            d = None
        subs = []
        ## open the file for reading
        with open(i) as f:
            ## parse the srt file
            subs = list(srt.parse(f))
            ## build only if we want multi-file
            if not values.settings.single_file:
                d = createDocument(values)

            buildTable(values, subs, Path(i).stem, d)

            if not values.settings.single_file:
                closeDocument(values, Path(i).stem, version, d)

    if values.settings.single_file:
        p = Path(files[0])
        n = os.path.basename(os.path.normpath(os.path.dirname(p.absolute())))
        closeDocument(values, n, version, d)


def init(version: str) -> dict:
    """Load settings from yaml"""

    ## so this beast gets the path to the settings file as an absolute path
    ## to the running script
    p = os.path.join(
        os.path.dirname(os.path.realpath(inspect.stack()[0][1])),
        "srt2docx_settings.yaml",
    )

    ## load settings
    values = munchify(yaml.safe_load(open(p)))

    ## do the work
    files = readFiles(values)
    processFiles(values, files, version)

    return values
